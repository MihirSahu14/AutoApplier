"""Render tailored resume + cover letter to DOCX.

Single-column professional resume layout in Calibri 10pt with bold section
headings and a thin bottom border.  Cover letter is a clean letter format with
the same header so the package looks unified.
"""
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches, RGBColor


# ---------- low-level helpers ----------

def _set_run(run, *, font="Calibri", size=10, bold=False, italic=False, color=None):
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._element.rPr
    if rPr is not None:
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        rFonts.set(qn("w:ascii"), font)
        rFonts.set(qn("w:hAnsi"), font)


def _no_space(p, *, before=0, after=0, line=1.0):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def _bottom_border(p):
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _section_heading(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    _set_run(run, size=11, bold=True)
    _no_space(p, before=6, after=2)
    _bottom_border(p)


def _two_col_line(doc, left_runs, right_runs):
    """A row with `left_runs` flush left and `right_runs` flush right.
    Each runs entry is (text, {style kwargs})."""
    p = doc.add_paragraph()
    _no_space(p, after=0)
    for text, style in left_runs:
        _set_run(p.add_run(text), **style)
    p.add_run("\t")
    # add right-aligned tab stop near right margin
    tabs = p.paragraph_format.tab_stops
    tabs.add_tab_stop(Inches(7.0), alignment=2)  # 2 = RIGHT
    for text, style in right_runs:
        _set_run(p.add_run(text), **style)
    return p


def _bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.05)
    _no_space(p, after=0)
    run = p.runs[0] if p.runs else p.add_run(text)
    if not p.runs or p.runs[0].text == "":
        run = p.add_run(text)
    else:
        run.text = text
    _set_run(run, size=10)


def _set_margins(doc, inches=0.5):
    for section in doc.sections:
        section.top_margin = Inches(inches)
        section.bottom_margin = Inches(inches)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)


# ---------- header ----------

def _header(doc, hdr: dict):
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _no_space(name_p, after=0)
    _set_run(name_p.add_run(hdr.get("name", "")), size=18, bold=True)

    pieces = []
    if hdr.get("location"): pieces.append(hdr["location"])
    if hdr.get("phone"):    pieces.append(hdr["phone"])
    if hdr.get("email"):    pieces.append(hdr["email"])
    for link in hdr.get("links", []):
        pieces.append(link.get("url") or link.get("label", ""))
    contact_p = doc.add_paragraph(" | ".join(filter(None, pieces)))
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _no_space(contact_p, after=2)
    for r in contact_p.runs:
        _set_run(r, size=9)


# ---------- public renderers ----------

def render_resume(data: dict, out_path: Path):
    doc = Document()
    _set_margins(doc)

    # base style
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)

    _header(doc, data.get("header", {}))

    # Education
    if data.get("education"):
        _section_heading(doc, "Education")
        for ed in data["education"]:
            _two_col_line(
                doc,
                left_runs=[(ed.get("school", ""), {"size": 10, "bold": True})],
                right_runs=[(ed.get("location", ""), {"size": 10})],
            )
            _two_col_line(
                doc,
                left_runs=[(ed.get("degree", ""), {"size": 10, "italic": True})],
                right_runs=[(ed.get("dates", ""), {"size": 10, "italic": True})],
            )
            for b in ed.get("bullets", []):
                _bullet(doc, b)

    # Experience
    if data.get("experience"):
        _section_heading(doc, "Experience")
        for exp in data["experience"]:
            _two_col_line(
                doc,
                left_runs=[(exp.get("company", ""), {"size": 10, "bold": True})],
                right_runs=[(exp.get("location", ""), {"size": 10})],
            )
            _two_col_line(
                doc,
                left_runs=[(exp.get("title", ""), {"size": 10, "italic": True})],
                right_runs=[(exp.get("dates", ""), {"size": 10, "italic": True})],
            )
            for b in exp.get("bullets", []):
                _bullet(doc, b)

    # Projects
    if data.get("projects"):
        _section_heading(doc, "Projects")
        for prj in data["projects"]:
            left = []
            left.append((prj.get("name", ""), {"size": 10, "bold": True}))
            tech = prj.get("tech", "")
            url = prj.get("url", "")
            extra = " — " + tech if tech else ""
            if url:
                extra += f" | {url}"
            if extra:
                left.append((extra, {"size": 10, "italic": True}))
            _two_col_line(
                doc,
                left_runs=left,
                right_runs=[(prj.get("date", ""), {"size": 10, "italic": True})],
            )
            for b in prj.get("bullets", []):
                _bullet(doc, b)

    # Skills
    if data.get("skills"):
        _section_heading(doc, "Technical Skills")
        for category, items in data["skills"].items():
            p = doc.add_paragraph()
            _no_space(p, after=0)
            _set_run(p.add_run(f"{category}: "), size=10, bold=True)
            _set_run(p.add_run(", ".join(items)), size=10)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    return out_path


def render_cover_letter(letter_text: str, header: dict, company: str, out_path: Path):
    doc = Document()
    _set_margins(doc)
    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)

    _header(doc, header)

    # date
    p = doc.add_paragraph(date.today().strftime("%B %d, %Y"))
    _no_space(p, before=10, after=10)

    # recipient line
    p = doc.add_paragraph(f"Hiring Team\n{company}")
    _no_space(p, after=14)

    # body
    for para in [s.strip() for s in letter_text.split("\n\n") if s.strip()]:
        body = doc.add_paragraph(para)
        body.paragraph_format.space_after = Pt(8)
        body.paragraph_format.line_spacing = 1.15

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    return out_path
